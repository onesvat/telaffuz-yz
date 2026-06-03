#!/usr/bin/env python3
"""Build the İÜC submission ``thesis/thesis.docx`` from the MyST chapters.

MyST's native Word export is per-article, so it cannot emit a single combined
thesis document. This script instead concatenates the chapters in TOC order,
rewrites the few MyST-specific constructs that plain Pandoc does not understand
(``{cite:p}`` / ``{cite:t}`` roles, ``{numref}`` references, and ``{figure}`` /
``{table}`` directives), assigns sequential figure/table numbers, expands pipe
tables into row-preserving list blocks for stable DOCX/PDF layout, and runs
Pandoc with its built-in citeproc against ``refs.bib``. Standard Markdown math
(``$$``) passes through to Pandoc natively.

The body is rendered against the İÜ-Cerrahpaşa style template
(``guides/template-tr.docx``) as Pandoc's ``--reference-doc``, which carries the
İÜ-C document defaults (Times New Roman 12 pt, A4 margins 3/2/3/2.5 cm, the
Balk1-9 heading styles). The İÜ-C cover page (``guides/cover-tr.docx``) is then
prepended with docxcompose; the title and date are filled in, while the
candidate-, advisor- and department-specific fields stay as template
placeholders to be completed in Word.

Run::

    uv run python scripts/build_thesis_docx.py
"""
from __future__ import annotations

import re
import subprocess
import sys
import zipfile
from pathlib import Path

import docx
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docxcompose.composer import Composer

REPO_ROOT = Path(__file__).resolve().parent.parent
THESIS = REPO_ROOT / "thesis"
CHAPTERS = THESIS / "chapters"
GUIDES = THESIS / "guides"
REFS = THESIS / "refs.bib"
OUT_DOCX = THESIS / "thesis.docx"
TEMPLATE_DOCX = GUIDES / "template-tr.docx"
COVER_DOCX = GUIDES / "cover-tr.docx"

TITLE = "Yapay Zeka Tabanlı Türkçe Diksiyon Değerlendirme Sistemi"
SUBTITLE = "Fonem Düzeyi Geri Bildirim ve Pedagojik Güvenlik"
COVER_TITLE = f"{TITLE}: {SUBTITLE}"
COVER_DATE = "Mayıs, 2026"

# İÜ-C cover placeholders (text-box runs in word/document.xml). Only the
# non-personal fields are filled; candidate/advisor/department are left for the
# author to complete in Word.
COVER_REPLACEMENTS = {
    "Tez Adını Giriniz": COVER_TITLE,
    "Kasım, 2021": COVER_DATE,
}

# İÜ-C numbers body chapters (1. Giriş … 6. Sonuç) but leaves front/back matter
# unnumbered. The template's Heading 1 (Balk1) auto-numbers, so these sections
# are reassigned to the template's unnumbered-heading style instead.
UNNUMBERED_HEADING_STYLE = "Başlık (Numarasız-Ortalı)"
UNNUMBERED_SECTIONS = {"Özet", "Abstract", "Kısaltmalar", "Kaynaklar"}
BODY_SECTIONS = {
    "Giriş",
    "Kavramsal Çerçeve",
    "Yöntem",
    "Bulgular",
    "Tartışma",
    "Sonuç ve Öneriler",
}
BODY_SECTION_ORDER = [
    "Giriş",
    "Kavramsal Çerçeve",
    "Yöntem",
    "Bulgular",
    "Tartışma",
    "Sonuç ve Öneriler",
]
BODY_SECTION_NUMBERS = {
    title: index + 1 for index, title in enumerate(BODY_SECTION_ORDER)
}
STYLE_REMAP = {
    "CaptionedFigure": "Normal",
    "FirstParagraph": "Body Text",
    "Compact": "Body Text",
}
AUTO_NUMBERED_HEADING_STYLE_IDS = (
    "Heading1",
    "Heading2",
    "Heading3",
    "Balk1",
    "Balk2",
    "Balk3",
)

# Body order (kaynaklar.md is only a {bibliography} directive; citeproc emits
# the reference list into the trailing #refs div instead).
CHAPTER_ORDER = [
    "ozet",
    "kisaltmalar",
    "giris",
    "kavramsal-cerceve",
    "yontem",
    "bulgular",
    "tartisma",
    "sonuc",
]

_CITE_RE = re.compile(r"\{cite:[pt]\}`([^`]+)`")
_NUMREF_RE = re.compile(r"\{numref\}`([^`]+)`(?P<suffix>'(?:d[ae]|t[ae]))?")
_FIGURE_RE = re.compile(
    r"^```\{figure\}\s*(?P<path>\S+)\s*\n"
    r"(?P<opts>(?::[^\n]*\n)*)"
    r"\s*\n"
    r"(?P<caption>.*?)\n"
    r"```\s*$",
    re.MULTILINE | re.DOTALL,
)
_TABLE_RE = re.compile(
    r"^```\{table\}\s*(?P<caption>[^\n]*)\n"
    r"(?P<opts>(?::[^\n]*\n)*)"
    r"\s*\n"
    r"(?P<body>.*?)\n"
    r"```\s*$",
    re.MULTILINE | re.DOTALL,
)
_HEADING_RE = re.compile(r"^(?P<marks>#{1,3})\s+(?P<title>.+?)\s*$")
_PIPE_TABLE_SEPARATOR_RE = re.compile(r"^:?-{3,}:?$")

# A pipe table immediately preceded by this marker is kept as a real grid table
# (numeric result tables read better aligned) instead of being expanded into a
# list block; the marker line itself is dropped from the output. The table is
# collected and replaced with a placeholder paragraph that is later swapped for a
# natively built grid (Pandoc's own Word tables render unreliably to PDF).
GRID_TABLE_MARKER = "<!-- grid-table -->"
GRID_TABLE_PLACEHOLDER = "GRIDTABLEPLACEHOLDER"


def _cite_replace(match: re.Match[str]) -> str:
    keys = [k.strip() for k in match.group(1).split(",") if k.strip()]
    return "[" + "; ".join(f"@{k}" for k in keys) + "]"


def collect_figure_numbers(text: str) -> dict[str, int]:
    """Map each figure ``:name:`` label to a sequential figure number."""
    numbers: dict[str, int] = {}
    n = 0
    for match in _FIGURE_RE.finditer(text):
        n += 1
        name = re.search(r":name:\s*(\S+)", match.group("opts"))
        if name:
            numbers[name.group(1)] = n
    return numbers


def collect_table_numbers(text: str) -> dict[str, int]:
    """Map each MyST table ``:name:`` label to a sequential table number."""
    numbers: dict[str, int] = {}
    n = 0
    for match in _TABLE_RE.finditer(text):
        n += 1
        name = re.search(r":name:\s*(\S+)", match.group("opts"))
        if name:
            numbers[name.group(1)] = n
    return numbers


def turkish_locative_suffix(number: int) -> str:
    """Return the apostrophe + locative suffix for a visible figure/table number."""
    if number in {1, 2, 7, 8, 11, 12}:
        return "'de"
    if number in {3, 4, 5}:
        return "'te"
    if number in {6, 9, 10}:
        return "'da"
    return "'de"


def rewrite(
    text: str, fig_numbers: dict[str, int], table_numbers: dict[str, int]
) -> str:
    def figure_replace(match: re.Match[str]) -> str:
        opts = match.group("opts")
        name = re.search(r":name:\s*(\S+)", opts)
        width = re.search(r":width:\s*(\S+)", opts)
        number = fig_numbers.get(name.group(1)) if name else None
        caption = " ".join(match.group("caption").split())
        prefix = f"**Şekil {number}.** " if number else ""
        path = match.group("path")
        # Chapter paths are ../../figures/... relative to chapters/; resolve to
        # an absolute path so Pandoc finds the image regardless of CWD.
        abs_path = (CHAPTERS / path).resolve()
        attr = f"{{width={width.group(1)}}}" if width else ""
        # Trailing newline guarantees a blank line before any heading that
        # immediately follows the figure block, so Pandoc keeps the heading a
        # heading instead of folding "## …" into the image paragraph.
        return f"![{prefix}{caption}]({abs_path}){attr}\n"

    def table_replace(match: re.Match[str]) -> str:
        opts = match.group("opts")
        name = re.search(r":name:\s*(\S+)", opts)
        number = table_numbers.get(name.group(1)) if name else None
        caption = " ".join(match.group("caption").split())
        prefix = f"**Tablo {number}.** " if number else ""
        body = match.group("body").strip()
        return f"{prefix}{caption}\n\n{body}\n"

    text = _FIGURE_RE.sub(figure_replace, text)
    text = _TABLE_RE.sub(table_replace, text)

    def numref_replace(match: re.Match[str]) -> str:
        label = match.group(1)
        suffix = match.group("suffix")
        if label in fig_numbers:
            number = fig_numbers[label]
            return f"Şekil {number}{turkish_locative_suffix(number) if suffix else ''}"
        if label in table_numbers:
            number = table_numbers[label]
            return f"Tablo {number}{turkish_locative_suffix(number) if suffix else ''}"
        return label + (suffix or "")

    text = _NUMREF_RE.sub(numref_replace, text)
    text = _CITE_RE.sub(_cite_replace, text)
    return text


def build_source() -> tuple[str, list[dict]]:
    raw = "\n\n".join(
        (CHAPTERS / f"{name}.md").read_text(encoding="utf-8")
        for name in CHAPTER_ORDER
    )
    raw = add_manual_heading_numbers(raw)
    fig_numbers = collect_figure_numbers(raw)
    table_numbers = collect_table_numbers(raw)
    body = rewrite(raw, fig_numbers, table_numbers)
    grid_tables: list[dict] = []
    body = rewrite_pipe_tables_to_lists(body, grid_tables)
    front = (
        "---\n"
        "lang: tr\n"
        "---\n\n"
    )
    refs = "\n\n# Kaynaklar\n\n::: {#refs}\n:::\n"
    return front + body + refs, grid_tables


def is_pipe_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def split_pipe_table_line(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def is_pipe_table_separator(line: str) -> bool:
    if not is_pipe_table_line(line):
        return False
    cells = split_pipe_table_line(line)
    return bool(cells) and all(_PIPE_TABLE_SEPARATOR_RE.match(cell) for cell in cells)


def pipe_table_to_list(header: list[str], rows: list[list[str]]) -> list[str]:
    """Convert every table row to a list item without dropping sparse cells."""
    block: list[str] = []
    width = max(len(header), *(len(row) for row in rows)) if rows else len(header)
    headers = [*header, *[""] * (width - len(header))]
    for row in rows:
        cells = [*row, *[""] * (width - len(row))]
        label = cells[0] or "Satır"
        if width == 2:
            block.append(f"- **{label}:** {cells[1]}")
            continue
        details = []
        for title, value in zip(headers[1:], cells[1:]):
            if value:
                details.append(f"{title or 'Değer'}: {value}")
        suffix = "; ".join(details)
        block.append(f"- **{label}:** {suffix}" if suffix else f"- **{label}**")
    return block


def column_alignments(separator_cells: list[str]) -> list[str]:
    """Map a pipe-table separator row to per-column alignment names."""
    aligns: list[str] = []
    for cell in separator_cells:
        cell = cell.strip()
        left = cell.startswith(":")
        right = cell.endswith(":")
        if left and right:
            aligns.append("center")
        elif right:
            aligns.append("right")
        else:
            aligns.append("left")
    return aligns


def rewrite_pipe_tables_to_lists(text: str, grid_tables: list[dict]) -> str:
    """Render Markdown pipe tables as stable, row-preserving list blocks.

    Pandoc's own Word tables render unreliably through the DOCX-to-PDF path, so
    tables are expanded into list blocks by default. A table preceded by the
    ``<!-- grid-table -->`` marker is instead collected into ``grid_tables`` and
    replaced with a placeholder paragraph; ``insert_grid_tables`` later swaps that
    placeholder for a natively built grid (numeric results read better aligned).
    """
    lines = text.splitlines()
    out: list[str] = []
    i = 0
    keep_as_grid = False
    while i < len(lines):
        if lines[i].strip() == GRID_TABLE_MARKER:
            keep_as_grid = True
            i += 1
            continue
        if (
            i + 1 < len(lines)
            and is_pipe_table_line(lines[i])
            and is_pipe_table_separator(lines[i + 1])
        ):
            header = split_pipe_table_line(lines[i])
            aligns = column_alignments(split_pipe_table_line(lines[i + 1]))
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and is_pipe_table_line(lines[i]):
                rows.append(split_pipe_table_line(lines[i]))
                i += 1
            if keep_as_grid:
                index = len(grid_tables)
                grid_tables.append(
                    {"header": header, "rows": rows, "aligns": aligns}
                )
                out.append(f"{GRID_TABLE_PLACEHOLDER}{index}")
                keep_as_grid = False
            else:
                out.extend(pipe_table_to_list(header, rows))
            continue
        out.append(lines[i])
        i += 1
    return "\n".join(out)


def add_manual_heading_numbers(text: str) -> str:
    """Number body headings in source text instead of relying on Word lists."""
    lines: list[str] = []
    current_chapter = 0
    current_section = 0
    current_subsection = 0
    for line in text.splitlines():
        match = _HEADING_RE.match(line)
        if not match:
            lines.append(line)
            continue

        marks = match.group("marks")
        title = match.group("title")
        if len(marks) == 1:
            number = BODY_SECTION_NUMBERS.get(title)
            if number is None:
                current_chapter = 0
                current_section = 0
                current_subsection = 0
                lines.append(line)
                continue
            current_chapter = number
            current_section = 0
            current_subsection = 0
            lines.append(f"# {current_chapter}. {title}")
            continue

        if current_chapter == 0:
            lines.append(line)
            continue

        if len(marks) == 2:
            current_section += 1
            current_subsection = 0
            lines.append(f"## {current_chapter}.{current_section}. {title}")
            continue

        current_subsection += 1
        if current_section == 0:
            lines.append(line)
        else:
            lines.append(
                f"### {current_chapter}.{current_section}.{current_subsection}. {title}"
            )
    return "\n".join(lines)


def build_filled_cover(dest: Path) -> None:
    """Copy the İÜ-C cover and fill the title/date text-box runs.

    The placeholders live inside positioned text boxes, so they are rewritten by
    string substitution on ``word/document.xml`` rather than via python-docx
    paragraphs. Each target string is contiguous within a single run.
    """
    with zipfile.ZipFile(COVER_DOCX) as zin:
        items = {name: zin.read(name) for name in zin.namelist()}
    xml = items["word/document.xml"].decode("utf-8")
    for placeholder, value in COVER_REPLACEMENTS.items():
        xml = xml.replace(placeholder, value)
    items["word/document.xml"] = xml.encode("utf-8")
    with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in items.items():
            zout.writestr(name, data)


def plain_heading_text(text: str) -> str:
    """Return a heading title without manual numeric prefixes."""
    return re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", text.strip())


def apply_iuc_heading_numbering(document: docx.Document) -> None:
    """Leave front/back matter unnumbered so body chapters number from Giriş = 1.

    The TOC heading and the Özet/Abstract/Kısaltmalar/Kaynaklar level-1 headings
    are switched from the auto-numbered Heading 1 to the template's unnumbered
    centred heading style.
    """
    unnumbered = document.styles[UNNUMBERED_HEADING_STYLE]
    normal = document.styles["Normal"]
    for paragraph in document.paragraphs:
        if paragraph.style is None:
            continue
        if paragraph.text.strip() in {"Table of Contents", "Contents"}:
            paragraph.text = "İçindekiler"
        style_name = paragraph.style.name
        if style_name in STYLE_REMAP:
            paragraph.style = document.styles[STYLE_REMAP[style_name]]
            style_name = paragraph.style.name
        if style_name == "ImageCaption":
            paragraph.style = normal
            for run in paragraph.runs:
                run.italic = True
                run.font.size = Pt(10)
            continue
        if style_name == "TOC Heading" or (
            style_name == "Heading 1"
            and plain_heading_text(paragraph.text) in UNNUMBERED_SECTIONS
        ):
            paragraph.style = unnumbered
        if style_name == "Heading 1" and plain_heading_text(paragraph.text) in BODY_SECTIONS:
            paragraph.paragraph_format.page_break_before = True


_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
}


def build_grid_table(document: docx.Document, spec: dict):
    """Build a native İÜ-C grid table element from a collected pipe table.

    Pandoc's own Word tables collapse when exported to PDF, so the few numeric
    result tables are rebuilt here with python-docx against the template's
    ``Table Grid`` style: bordered cells, a bold header row, per-column alignment
    from the Markdown separator, and a reduced font for wider tables. The table is
    detached from the document tail and returned for placement at the marker.
    """
    header = spec["header"]
    rows = spec["rows"]
    aligns = spec["aligns"]
    ncols = len(header)
    font_size = Pt(9.0 if ncols >= 5 else 10.0)

    table = document.add_table(rows=1 + len(rows), cols=ncols)
    try:
        table.style = document.styles["Table Grid"]
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    def fill(cells, values, *, bold):
        for col, cell in enumerate(cells):
            value = values[col] if col < len(values) else ""
            align = aligns[col] if col < len(aligns) else "left"
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.bold = bold
            run.font.size = font_size

    fill(table.rows[0].cells, header, bold=True)
    for row_index, row in enumerate(rows, start=1):
        fill(table.rows[row_index].cells, row, bold=False)

    tbl = table._tbl
    tbl.getparent().remove(tbl)
    return tbl


def insert_grid_tables(document: docx.Document, grid_tables: list[dict]) -> None:
    """Replace grid-table placeholder paragraphs with native grid tables."""
    if not grid_tables:
        return
    wanted = {
        f"{GRID_TABLE_PLACEHOLDER}{index}": spec
        for index, spec in enumerate(grid_tables)
    }
    for paragraph in list(document.paragraphs):
        spec = wanted.get(paragraph.text.strip())
        if spec is None:
            continue
        tbl = build_grid_table(document, spec)
        paragraph._p.addprevious(tbl)
        paragraph._p.getparent().remove(paragraph._p)


def style_code_blocks(document: docx.Document) -> None:
    """Render fenced pseudocode as a clean, bordered monospace box.

    Pandoc maps fenced code to the empty ``Source Code`` paragraph style, which
    otherwise inherits the justified Times New Roman body font: proportional
    spacing breaks the column alignment of the algorithm listings and long lines
    overflow the İÜ-C text block. This restyles the listings as a left-aligned
    DejaVu Sans Mono box (light fill, thin border, kept together on one page) and
    keeps each ``Algoritma N.`` caption attached to the box that follows it.
    """
    try:
        style = document.styles["Source Code"]
    except KeyError:
        return

    style.font.name = "DejaVu Sans Mono"
    style.font.size = Pt(9)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "DejaVu Sans Mono")

    ppr = style.element.get_or_add_pPr()
    for child in list(ppr):
        ppr.remove(child)
    # CT_PPr child order: keepLines, pBdr, shd, spacing, ind, jc.
    ppr.append(OxmlElement("w:keepLines"))

    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")  # quarter-points → 0.5 pt rule
        border.set(qn("w:space"), "6")  # pt of padding between rule and text
        border.set(qn("w:color"), "AAAAAA")
        pbdr.append(border)
    ppr.append(pbdr)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    ppr.append(shd)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:after"), "120")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "170")
    ind.set(qn("w:right"), "170")
    ppr.append(ind)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    ppr.append(jc)

    caption_re = re.compile(r"^Algoritma\s+\d+\.")
    for paragraph in document.paragraphs:
        if caption_re.match(paragraph.text.strip()):
            paragraph.paragraph_format.keep_with_next = True


def remove_numpr_from_style(xml: str, style_id: str) -> str:
    """Remove Word auto-numbering from a paragraph style definition."""
    pattern = re.compile(
        rf'(<w:style\b(?=[^>]*w:styleId="{re.escape(style_id)}")[\s\S]*?</w:style>)'
    )

    def replace(match: re.Match[str]) -> str:
        return re.sub(r"<w:numPr>[\s\S]*?</w:numPr>", "", match.group(1))

    return pattern.sub(replace, xml, count=1)


def remove_numpr_from_heading_paragraphs(xml: str) -> str:
    """Remove paragraph-level auto-numbering from generated heading paragraphs."""
    style_ids = "|".join((*AUTO_NUMBERED_HEADING_STYLE_IDS, "TOCHeading"))
    pattern = re.compile(
        rf'(<w:pPr>(?:(?!</w:pPr>).)*<w:pStyle w:val="(?:{style_ids})"/>'
        rf'(?:(?!</w:pPr>).)*)(<w:numPr>[\s\S]*?</w:numPr>)'
        rf'((?:(?!</w:pPr>).)*</w:pPr>)'
    )
    previous = None
    while previous != xml:
        previous = xml
        xml = pattern.sub(r"\1\3", xml)
    return xml


def patch_generated_docx_xml(docx_path: Path) -> None:
    """Fix generated DOCX XML that Pandoc/python-docx cannot fully expose.

    ``python-docx`` does not expose the generated TOC heading because Pandoc
    places it under ``w:sdt``. Without this patch the PDF export shows the
    English heading. Heading numbering is rendered as explicit text in the
    Markdown source, so Word's automatic heading counters are stripped here.
    """
    tmp_path = docx_path.with_suffix(".docx.tmp")
    old_heading = (
        '<w:p><w:pPr><w:pStyle w:val="TBal"/></w:pPr><w:r>'
        '<w:t xml:space="preserve">Table of Contents</w:t></w:r></w:p>'
    )
    new_heading = (
        '<w:p><w:pPr><w:pStyle w:val="TOCHeading"/></w:pPr><w:r>'
        '<w:t xml:space="preserve">İçindekiler</w:t></w:r></w:p>'
    )
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
        tmp_path, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                xml = data.decode("utf-8")
                xml = xml.replace(
                    '<w:docPartGallery w:val="Table of Contents"/>',
                    '<w:docPartGallery w:val="İçindekiler"/>',
                    1,
                )
                if old_heading in xml:
                    xml = xml.replace(old_heading, new_heading, 1)
                xml = remove_numpr_from_heading_paragraphs(xml)
                data = xml.encode("utf-8")
            elif item.filename == "word/styles.xml":
                xml = data.decode("utf-8")
                for style_id in AUTO_NUMBERED_HEADING_STYLE_IDS:
                    xml = remove_numpr_from_style(xml, style_id)
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    tmp_path.replace(docx_path)


def main() -> None:
    if not REFS.exists():
        sys.exit(f"missing bibliography: {REFS}")
    source, grid_tables = build_source()
    build_dir = THESIS / "_build"
    build_dir.mkdir(parents=True, exist_ok=True)
    combined = build_dir / "thesis-combined.md"
    combined.write_text(source, encoding="utf-8")

    body_docx = build_dir / "thesis-body.docx"
    cmd = [
        "pandoc",
        str(combined),
        "--from=markdown+tex_math_dollars",
        "--citeproc",
        f"--bibliography={REFS}",
        f"--reference-doc={TEMPLATE_DOCX}",
        "--toc",
        "--toc-depth=2",
        "-o",
        str(body_docx),
    ]
    subprocess.run(cmd, check=True, cwd=THESIS)

    # Prepend the İÜ-C cover (title/date filled, personal fields left blank).
    cover_filled = build_dir / "cover-filled.docx"
    build_filled_cover(cover_filled)
    master = docx.Document(str(cover_filled))
    master.add_page_break()
    Composer(master).append(docx.Document(str(body_docx)))
    apply_iuc_heading_numbering(master)
    insert_grid_tables(master, grid_tables)
    style_code_blocks(master)
    # The cover (absolutely positioned text boxes) is margin-independent, so the
    # İÜ-C body margins (left 3, right 2, top 3, bottom 2.5 cm) are applied to
    # every section without affecting the cover layout.
    for section in master.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2.5)
    master.save(str(OUT_DOCX))
    patch_generated_docx_xml(OUT_DOCX)
    print(f"OK: {OUT_DOCX.relative_to(REPO_ROOT)}")


if __name__ == "__main__":
    main()
